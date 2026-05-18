"""Build a finished CV docx by replacing the content of each section in the template.

Strategy:
  1. Open template/cv-template.docx
  2. For each leaf section:
       a. Capture each existing paragraph's <w:pPr> (paragraph properties: indent,
          tab stops, line spacing). These define the section's visual layout.
       b. Capture the first run's <w:rPr> as a font template (font face, size).
       c. Delete the existing content paragraphs (but NOT <w:sectPr> or any
          other body child — only <w:p> elements).
       d. Insert new paragraphs whose pPr is cloned positionally from the
          original (with cyclic repeat if new content has more paragraphs).
          Each new paragraph carries the section's font rPr.
  3. For peer_reviewed_original, the new paragraphs are publication citations
     emitted as multi-run rich text so the author name can be bolded.
  4. Rewrite header text in case the {count} placeholder needs updating.
"""

import copy
import json
import re
import sys
from pathlib import Path

import yaml
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT / "src"))

from anchors import find_anchors
from sections import SECTIONS
from pubmed import format_citation


# ----- low-level XML helpers -----

def _t_element(text):
    el = OxmlElement("w:t")
    el.text = text
    if text != text.strip() or "  " in text:
        el.set(qn("xml:space"), "preserve")
    return el


def _tab_element():
    return OxmlElement("w:tab")


def _new_p():
    return OxmlElement("w:p")


def _new_r(rpr_template=None):
    r = OxmlElement("w:r")
    if rpr_template is not None:
        r.insert(0, copy.deepcopy(rpr_template))
    return r


def _make_run(text, rpr_template=None, bold=False):
    """Create a <w:r> with the given text, cloning the font properties from
    rpr_template, and optionally setting bold."""
    r = _new_r(rpr_template)
    if bold:
        rpr = r.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            r.insert(0, rpr)
        # Remove any existing w:b; add new one
        for b in rpr.findall(qn("w:b")):
            rpr.remove(b)
        rpr.append(OxmlElement("w:b"))
    if "\t" in text:
        bits = text.split("\t")
        for i, bit in enumerate(bits):
            if i > 0:
                r.append(_tab_element())
            if bit:
                r.append(_t_element(bit))
    else:
        r.append(_t_element(text))
    return r


# ----- section template capture -----

_TOGGLE_TAGS = {qn(t) for t in ("w:b", "w:bCs", "w:i", "w:iCs", "w:u", "w:strike",
                                 "w:dstrike", "w:caps", "w:smallCaps", "w:vanish",
                                 "w:webHidden", "w:highlight", "w:shd")}


def _clean_rpr(rpr):
    """Strip dynamic toggle attributes (bold/italic/highlight/etc.) from a copied
    rPr so it only carries font identity (face, size, color)."""
    rpr = copy.deepcopy(rpr)
    for child in list(rpr):
        if child.tag in _TOGGLE_TAGS:
            rpr.remove(child)
    return rpr


def capture_section_template(doc, header_idx, end_idx):
    """Return (pprs, rpr_template) for the paragraphs in (header_idx, end_idx).

    pprs is a list of deep-copied <w:pPr> elements (or None for paragraphs that
    had no pPr), one per paragraph in the original content. rpr_template carries
    font identity only (toggle attributes like bold/italic stripped — those are
    applied per-run as needed).
    """
    pprs = []
    rpr_template = None
    for i in range(header_idx + 1, end_idx):
        p = doc.paragraphs[i]
        ppr = p._element.find(qn("w:pPr"))
        pprs.append(copy.deepcopy(ppr) if ppr is not None else None)
        if rpr_template is None:
            for r in p._element.findall(qn("w:r")):
                rpr = r.find(qn("w:rPr"))
                if rpr is not None and len(rpr) > 0:
                    rpr_template = _clean_rpr(rpr)
                    if len(rpr_template) == 0:
                        rpr_template = None  # only toggles, no font info
                    else:
                        break
    return pprs, rpr_template


def pick_ppr(pprs, position):
    """Pick a pPr template for the paragraph at the given position.

    Cycles through the captured pprs so a section like "grant entry: 3 lines +
    blank" repeats its 4-paragraph pattern as new grants are added.
    """
    if not pprs:
        return None
    return pprs[position % len(pprs)]


# Sections whose pPr we replace with another section's, to keep the
# "date / content" layout consistent across the CV. The original CV used
# different indents in different sections (some 2", some 1.5", some no indent
# relying on default tab stops). Harmonize to residency_fellowship's 2" hanging
# indent — that's the layout the user has consistently approved.
#
# Two variants:
#  - FALLBACK_TEMPLATE_SECTIONS: only applied when the section's own template
#    is empty (sections empty in the original docx, e.g. trainees).
#  - HARMONIZE_TEMPLATE_SECTIONS: always applied, even when the section has
#    its own pPrs — used to normalize visually inconsistent sections.
FALLBACK_TEMPLATE_SECTIONS = {
    "trainees_postdoc":   "residency_fellowship",
    "trainees_graduate":  "residency_fellowship",
    "trainees_med":       "residency_fellowship",
    "trainees_undergrad": "residency_fellowship",
    "trainees_highschool": "residency_fellowship",
}
HARMONIZE_TEMPLATE_SECTIONS = {
    "university_admin_service":  "residency_fellowship",
    "service_professional_orgs": "residency_fellowship",
    "editorial_service":         "residency_fellowship",
    "ad_hoc_reviewer":           "residency_fellowship",
    "grant_reviewer":            "residency_fellowship",
    "community_service":         "residency_fellowship",
    "teaching_stanford":         "residency_fellowship",
    "teaching_pre_stanford":     "residency_fellowship",
    "board_certification":       "residency_fellowship",
}


def _pprs_look_empty(pprs):
    """True if every captured pPr is None or has no children (effectively blank)."""
    return all(p is None or len(p) == 0 for p in pprs)


# ----- run-level: deletion of existing paragraphs -----

def delete_section_paragraphs(doc, header_idx, end_idx):
    """Delete <w:p> and <w:tbl> elements strictly between header and the next anchor.

    Word tables ride next to paragraphs as siblings inside <w:body>; if we don't
    remove them, the original template's tables (e.g., trainee lists) accumulate
    on top of the newly-inserted sheet-driven paragraphs. Never touches
    <w:sectPr> or other body children."""
    body = doc.paragraphs[header_idx]._element.getparent()
    header_el = doc.paragraphs[header_idx]._element
    end_el = doc.paragraphs[end_idx]._element if end_idx < len(doc.paragraphs) else None

    p_tag = qn("w:p")
    tbl_tag = qn("w:tbl")
    to_delete = []
    walking = header_el.getnext()
    while walking is not None and walking is not end_el:
        if walking.tag in (p_tag, tbl_tag):
            to_delete.append(walking)
        walking = walking.getnext()
    for el in to_delete:
        body.remove(el)
    return header_el


# ----- insertion of new paragraphs -----

SUBHEADING_MARKER = "[[H]]"


def insert_paragraph_segments(header_el, segments_list, pprs, rpr_template, style_name):
    """Insert N paragraphs after header_el. segments_list[i] is either a string
    OR a list of (text, attrs) tuples for the i-th paragraph.

    Strings starting with SUBHEADING_MARKER are rendered as bold, indent-free
    sub-heading paragraphs (used for in-section group titles like
    "Courses" / "Tutorials and Lectures" within teaching_pre_stanford).
    """
    cursor = header_el
    for pos, seg in enumerate(segments_list):
        new_p = _new_p()

        # Detect sub-heading: string starting with the marker
        is_subheading = isinstance(seg, str) and seg.startswith(SUBHEADING_MARKER)
        if is_subheading:
            seg = seg[len(SUBHEADING_MARKER):]  # strip marker
            # Sub-headings get an empty pPr (no indent, no list numbering) plus
            # a leading blank-line gap above for visual breathing room.
            ppr_el = OxmlElement("w:pPr")
            new_p.append(ppr_el)
        else:
            # pPr: clone from template if available
            ppr = pick_ppr(pprs, pos)
            if ppr is not None:
                new_p.append(copy.deepcopy(ppr))
            else:
                ppr_el = OxmlElement("w:pPr")
                pstyle = OxmlElement("w:pStyle")
                pstyle.set(qn("w:val"), style_name)
                ppr_el.append(pstyle)
                new_p.append(ppr_el)

        # Build runs
        if isinstance(seg, str):
            if seg:
                new_p.append(_make_run(seg, rpr_template, bold=is_subheading))
        else:
            for text, attrs in seg:
                bold = bool(attrs.get("bold")) if attrs else False
                new_p.append(_make_run(text, rpr_template, bold=bold))

        cursor.addnext(new_p)
        cursor = new_p


# ----- header text update (preserves bold via existing run rPr) -----

def _set_header_text(p, text):
    runs = p._element.findall(qn("w:r"))
    if not runs:
        p.add_run(text)
        return
    first = runs[0]
    for t in first.findall(qn("w:t")):
        first.remove(t)
    for tab in first.findall(qn("w:tab")):
        first.remove(tab)
    first.append(_t_element(text))
    for r in runs[1:]:
        p._element.remove(r)


# ----- author name bolding in publication citations -----

# Match "Westover" alone or with initials immediately after (e.g. "Westover MB",
# "Westover, MB", "Westover BM"). Also handles the rare form where initials
# come before the name. Trailing punctuation (period, comma) is NOT included so
# punctuation stays in a non-bold run.
_NAME_RE = re.compile(
    r"\bWestover(?:,?\s+(?:[A-Z]{1,3}|[A-Z]\.[A-Z]\.?))?",
)


def citation_segments(rec, annotation=""):
    """Return list of (text, attrs) tuples for a publication citation.

    Splits the rendered citation around Westover-name matches so we can render
    the author name as a bold run.
    """
    flat = format_citation(rec, annotation=annotation)
    segments = []
    last = 0
    for m in _NAME_RE.finditer(flat):
        if m.start() > last:
            segments.append((flat[last : m.start()], {}))
        segments.append((flat[m.start() : m.end()], {"bold": True}))
        last = m.end()
    if last < len(flat):
        segments.append((flat[last:], {}))
    return segments or [(flat, {})]


# ----- pub-list builder -----

def build_pub_segments_list(pubs_json_path, annotations_yaml_path):
    """Return (segments_list, annotation_indices).

    segments_list is a list of paragraphs; each paragraph is a list of
    (text, attrs) tuples. annotation_indices is the list of paragraph
    positions that are annotation paragraphs (so the assembler can apply the
    Colorful List style to them later).
    """
    pubs = json.loads(Path(pubs_json_path).read_text())
    annotations = {}
    if Path(annotations_yaml_path).exists():
        annotations = yaml.safe_load(Path(annotations_yaml_path).read_text()) or {}

    paragraphs = []
    annotation_indices = []
    for pub in pubs:
        paragraphs.append(citation_segments(pub))
        notes = annotations.get(pub["pmid"], [])
        for note in notes:
            annotation_indices.append(len(paragraphs))
            paragraphs.append([(note, {})])
    return paragraphs, annotation_indices


# ----- entry point -----

def fill_count(template, count):
    return template.replace("{count}", str(count))


def assemble(out_path):
    template_path = ROOT / "template" / "cv-template.docx"
    sections_dir = ROOT / "cache" / "sections"

    doc = Document(template_path)

    section_meta = {s[0]: {"prefix": s[1], "header_template": s[2], "style": s[3], "source": s[4]}
                    for s in SECTIONS}

    # Build per-section content (list of paragraphs; each paragraph is either a
    # string or a list of (text, attrs) tuples).
    section_content = {}
    pubs_count = 0
    pub_annotation_indices = []

    for key, meta in section_meta.items():
        if meta["source"] == "ncbi":
            segs_list, ann_idx = build_pub_segments_list(
                ROOT / "cache" / "publications.json",
                ROOT / "manual_annotations.yaml",
            )
            section_content[key] = segs_list
            pubs_count = sum(1 for p in segs_list
                              if any("PMID:" in t or "PMCID:" in t
                                     for t, _attrs in (p if isinstance(p, list) else [(p, {})])))
            pub_annotation_indices = ann_idx
        else:
            fp = sections_dir / f"{key}.txt"
            text = fp.read_text() if fp.exists() else ""
            if text.endswith("\n"):
                text = text[:-1]
            lines = text.split("\n") if text else []
            if lines and lines[-1].strip():
                lines.append("")
            section_content[key] = lines

    # Process sections in reverse so deletions don't shift earlier indices.
    annotation_style = "Colorful List - Accent 11"
    for key in [s[0] for s in SECTIONS][::-1]:
        meta = section_meta[key]
        anchors = find_anchors(doc)
        my_idx = None
        next_idx = None
        for i, (idx, _t, kind) in enumerate(anchors):
            if kind == f"section:{key}":
                my_idx = idx
                next_idx = anchors[i + 1][0] if i + 1 < len(anchors) else len(doc.paragraphs)
                break
        if my_idx is None:
            continue

        # Capture pPr templates and font rPr from original section content.
        pprs, rpr_template = capture_section_template(doc, my_idx, next_idx)
        # Determine if this section borrows pPr from another section. Harmonize
        # entries take precedence and apply unconditionally; fallback only
        # applies if the section's own template is empty.
        borrow_key = None
        if key in HARMONIZE_TEMPLATE_SECTIONS:
            borrow_key = HARMONIZE_TEMPLATE_SECTIONS[key]
        elif _pprs_look_empty(pprs) and key in FALLBACK_TEMPLATE_SECTIONS:
            borrow_key = FALLBACK_TEMPLATE_SECTIONS[key]
        if borrow_key:
            for i, (idx, _t, kind) in enumerate(anchors):
                if kind == f"section:{borrow_key}":
                    fb_end = anchors[i + 1][0] if i + 1 < len(anchors) else len(doc.paragraphs)
                    fb_pprs, fb_rpr = capture_section_template(doc, idx, fb_end)
                    # Only use pPrs from CONTENT paragraphs (non-empty text in
                    # the borrowed section). This avoids accidentally cycling
                    # in a trailing artifact pPr (e.g., a 720/720 bold one) that
                    # snuck into the captured list. Use only one pPr — the first
                    # content-paragraph's — and reuse it for every new row, so
                    # the section is visually homogenous regardless of cycle
                    # length.
                    first_content_ppr = None
                    for j in range(idx + 1, fb_end):
                        p = doc.paragraphs[j]
                        if not p.text.strip():
                            continue
                        own_ppr = p._element.find(qn("w:pPr"))
                        if own_ppr is not None and len(own_ppr) > 0:
                            first_content_ppr = copy.deepcopy(own_ppr)
                            break
                    if first_content_ppr is not None:
                        pprs = [first_content_ppr]
                        if rpr_template is None:
                            rpr_template = fb_rpr
                    break

        # Delete the old content paragraphs.
        header_el = delete_section_paragraphs(doc, my_idx, next_idx)

        # Insert new content paragraphs.
        insert_paragraph_segments(
            header_el,
            section_content[key],
            pprs,
            rpr_template,
            meta["style"],
        )

        # Refresh header text with up-to-date counts.
        anchors = find_anchors(doc)
        for i, (idx, _t, kind) in enumerate(anchors):
            if kind == f"section:{key}":
                header_p = doc.paragraphs[idx]
                if "{count}" in meta["header_template"]:
                    if key == "peer_reviewed_original":
                        n = pubs_count
                    else:
                        n = sum(1 for ln in section_content[key]
                                if (ln if isinstance(ln, str) else "".join(t for t, _ in ln)).strip())
                    _set_header_text(header_p, fill_count(meta["header_template"], n))
                break

        # Apply Colorful List style to pub annotation paragraphs.
        if key == "peer_reviewed_original" and pub_annotation_indices:
            anchors = find_anchors(doc)
            for i, (idx, _t, kind) in enumerate(anchors):
                if kind == f"section:{key}":
                    pub_start = idx + 1
                    for offset in pub_annotation_indices:
                        ap = doc.paragraphs[pub_start + offset]
                        try:
                            ap.style = annotation_style
                        except KeyError:
                            pass
                    break

    # Update the "Curriculum Vitae: updated …" line at the top.
    from datetime import date
    today_str = date.today().strftime("%B %-d, %Y")
    if doc.paragraphs and doc.paragraphs[0].text.strip().startswith("Curriculum Vitae"):
        _set_header_text(doc.paragraphs[0], f"Curriculum Vitae: updated {today_str}")

    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    from datetime import date
    out = ROOT / "output" / f"1-Westover-SU-CV-{date.today().isoformat()}.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    assemble(out)
    print(f"Wrote {out}")
