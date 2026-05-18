"""One-shot: parse legacy free-form CV content (invited talks, trainees) and
push it into the Sheets created by bootstrap_via_drive.py.

Talks: parsed from the template docx (4 sections × ~20 talks each).
Trainees: parsed from "Formally Supervised Trainees.docx" (80 rows), then
classified into postdoc / grad / med / undergrad / highschool by inspecting
the description.

Re-runnable. Each section's existing data rows are REPLACED with the freshly
parsed rows; the header row is preserved.
"""

import csv
import io
import json
import re
import subprocess
import sys
from pathlib import Path

import requests
import yaml
from docx import Document

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT / "src"))

from anchors import section_ranges
from sheets_schema import SHEET_SCHEMAS


# ---------------------- talks ----------------------

US_STATES = (
    "AL|AK|AZ|AR|CA|CO|CT|DE|FL|GA|HI|ID|IL|IN|IA|KS|KY|LA|ME|MD|MA|MI|MN|"
    "MS|MO|MT|NE|NV|NH|NJ|NM|NY|NC|ND|OH|OK|OR|PA|RI|SC|SD|TN|TX|UT|VT|VA|"
    "WA|WV|WI|WY|DC"
)

TALK_RE = re.compile(
    r'^\s*[“"]\s*(?P<title>.+?)\.?\s*[”"]\s*'   # "Title."
    r'(?P<rest>.+?)\s*'                          # Venue, Location
    r'\((?P<year>\d{4})\)\s*\.?\s*$'             # (Year)
)


def parse_one_talk(line):
    m = TALK_RE.match(line)
    if not m:
        return None
    title = m.group("title").strip().rstrip(",.").strip()
    rest = m.group("rest").strip().rstrip(",")
    year = m.group("year")
    parts = [p.strip() for p in rest.split(",")]
    # If last two parts look like City, ST → Location = "City, ST"
    if len(parts) >= 2 and re.fullmatch(US_STATES, parts[-1]):
        venue = ", ".join(parts[:-2])
        location = ", ".join(parts[-2:])
    elif len(parts) >= 2:
        venue = ", ".join(parts[:-1])
        location = parts[-1]
    else:
        venue = ""
        location = parts[0] if parts else ""
    return {"Date": year, "Title": title, "Venue": venue, "Location": location}


def parse_talks_section(doc, key):
    """Return list of row-dicts from the named section in the template docx."""
    r = section_ranges(doc)
    _hdr, start, end = r[key]
    rows = []
    for i in range(start, end):
        t = doc.paragraphs[i].text.strip()
        if not t:
            continue
        parsed = parse_one_talk(t)
        if parsed:
            rows.append(parsed)
        else:
            # Couldn't parse — stuff the raw line into Venue so nothing is lost.
            rows.append({"Date": "", "Title": "", "Venue": t, "Location": ""})
    return rows


# ---------------------- trainees ----------------------

TRAINEES_SRC = Path("/Users/mwestover/Library/CloudStorage/Box-Box/Brandon - PHI/"
                    "!@@@-Work/CV/LessFrequentlyUsed/NSF/Formally Supervised Trainees.docx")


def classify_trainee(description):
    """Return the section key for this trainee description."""
    d = description.lower()
    # Most specific first
    if "high school" in d:
        return "trainees_highschool"
    if any(x in d for x in [
        "undergraduate", "senior at", "sophomore at", "freshman at", "junior at",
        " bs (", " bs,", "byu undergraduate", "mit undergraduate",
    ]):
        return "trainees_undergrad"
    if any(x in d for x in [
        "postdoctoral fellow", "postdoctoral research", "postdoc"
    ]):
        return "trainees_postdoc"
    # PhD-already + research role → postdoc
    if re.search(r"\bphd[,.]", d) and any(x in d for x in [
        "visiting scholar", "research fellow", "visiting researcher", "research assistant professor",
    ]):
        return "trainees_postdoc"
    if any(x in d for x in [
        "phd candidate", "phd student", "masters student", "msc", " ms (", " ms,",
        "mph candidate", "graduate student", "biomedical informatics program",
        "phd from",
    ]):
        return "trainees_graduate"
    # Default for clinical/medical trainees
    if any(x in d for x in [
        "md candidate", "md/phd candidate", "medical student", "md, phd", "mb, bch",
        "resident", "fellow", " md.", " md,", "harvard medical school /",
    ]):
        return "trainees_med"
    return "trainees_med"  # catch-all (most common)


CURRENT_RE = re.compile(r"\bcurrently\b[^.]*", re.IGNORECASE)


def extract_trainee_fields(dates_cell, desc_cell):
    """Return {Dates, Name, Role, Current Position} from a row of the docx table."""
    desc = re.sub(r"\s+", " ", desc_cell.strip())
    # Name is everything before the first comma or period (whichever comes first)
    name_match = re.match(r"^([^,.]+)([,.])\s*(.*)$", desc)
    if name_match:
        name = name_match.group(1).strip()
        rest = name_match.group(3).strip()
    else:
        name = desc.split()[0] if desc else ""
        rest = desc[len(name):].strip().lstrip(",.").strip()

    # Current position
    cur = ""
    m = CURRENT_RE.search(rest)
    if m:
        cur = m.group(0).strip().rstrip(".")
        # Strip leading "Currently"
        cur = re.sub(r"^currently\s+", "", cur, flags=re.IGNORECASE).strip()
        rest = rest[: m.start()].strip().rstrip(",.").strip()

    # Role: the first sentence-ish chunk after the name, capped at 120 chars
    role = rest
    # Trim long descriptions to the first sentence
    first_sent = re.split(r"(?<=[a-z])\.(?:\s+|$)", role, maxsplit=1)[0].strip()
    if len(first_sent) < len(role):
        role = first_sent
    role = role[:200].rstrip(",;")

    return {
        "Dates": (dates_cell or "").strip(),
        "Name": name,
        "Role": role,
        "Current Position": cur,
    }


def parse_trainees():
    """Return dict: section_key → list of row dicts."""
    d = Document(TRAINEES_SRC)
    out = {k: [] for k in [
        "trainees_postdoc", "trainees_graduate", "trainees_med",
        "trainees_undergrad", "trainees_highschool",
    ]}
    for row in d.tables[0].rows:
        cells = [c.text for c in row.cells]
        if len(cells) < 2:
            continue
        dates = cells[0].strip()
        desc = cells[1].strip()
        if not desc:
            continue
        bucket = classify_trainee(desc)
        out[bucket].append(extract_trainee_fields(dates, desc))
    return out


# ---------------------- upload via Drive + Sheets REST APIs ----------------------

DRIVE_API = "https://www.googleapis.com/drive/v3"
SHEETS_API = "https://sheets.googleapis.com/v4/spreadsheets"


def get_user_token():
    r = subprocess.run(
        ["gcloud", "auth", "print-access-token",
         "--scopes=https://www.googleapis.com/auth/drive"],
        capture_output=True, text=True, check=True,
    )
    return r.stdout.strip()


def replace_sheet_data(token, sheet_id, columns, rows):
    """Overwrite the file via Drive API uploadType=media with a fresh CSV.
    Drive auto-re-converts CSV → Sheet, preserving the sheet_id."""
    buf = io.StringIO()
    w = csv.writer(buf)
    w.writerow(columns)
    for row in rows:
        w.writerow([row.get(c, "") for c in columns])
    csv_bytes = buf.getvalue().encode("utf-8")
    r = requests.patch(
        f"https://www.googleapis.com/upload/drive/v3/files/{sheet_id}",
        params={"uploadType": "media"},
        headers={"Authorization": f"Bearer {token}",
                 "Content-Type": "text/csv"},
        data=csv_bytes, timeout=60,
    )
    r.raise_for_status()


# ---------------------- entry point ----------------------

def main():
    cfg = yaml.safe_load((ROOT / "config.yaml").read_text())
    section_cfg = cfg.get("sections") or {}
    token = get_user_token()

    # --- Talks ---
    doc = Document(ROOT / "template" / "cv-template.docx")
    talk_keys = ["invited_grand_rounds", "invited_local",
                 "invited_national", "invited_international"]
    for key in talk_keys:
        rows = parse_talks_section(doc, key)
        sid = (section_cfg.get(key) or {}).get("sheet_id")
        print(f"  {key}: parsed {len(rows)} rows → {sid[:10] if sid else '???'}...")
        if sid and rows:
            replace_sheet_data(token, sid, SHEET_SCHEMAS[key]["columns"], rows)

    # --- Trainees ---
    by_bucket = parse_trainees()
    for key, rows in by_bucket.items():
        sid = (section_cfg.get(key) or {}).get("sheet_id")
        print(f"  {key}: parsed {len(rows)} rows → {sid[:10] if sid else '???'}...")
        if sid:
            replace_sheet_data(token, sid, SHEET_SCHEMAS[key]["columns"], rows)


if __name__ == "__main__":
    main()
